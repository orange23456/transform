from docx import Document
from pathlib import Path

p = Path(r'C:\Users\hz-user\Desktop\加热搅拌系列产品.docx')
doc = Document(str(p))
paras = [x.text.strip() for x in doc.paragraphs if x.text.strip()]
print('paragraphs', len(paras))
print('tables', len(doc.tables))
print('first_paragraphs')
for x in paras[:60]:
    print(x[:300])
print('tables_preview')
for i, table in enumerate(doc.tables[:8], 1):
    print('TABLE', i, 'rows', len(table.rows), 'cols', len(table.columns))
    for row in table.rows[:6]:
        print(' | '.join(cell.text.replace('\n', ' / ')[:100] for cell in row.cells))
