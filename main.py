import argparse
import json
import os
import re
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.worksheet.table import Table, TableStyleInfo


HEADERS = [
    "来源",
    "商品中文名称",
    "型号",
    "产品名称（英文）",
    "URL",
    "简介",
    "产品详情",
    "代码",
    "网站分类",
    "主图",
    "详情图",
    "图片文件名",
    "alt text",
    "title",
    "Caption",
    "价格",
    "💲",
    "是否上传",
]

BLANK_UPLOAD_COLUMNS = {"来源", "是否上传", "alt text", "title", "Caption"}
SUPPORTED_EXTENSIONS = {".docx", ".xlsx", ".txt", ".md", ".csv"}
FORBIDDEN_PHRASES = [
    "separate from the family record",
    "separated from the family record",
    "separate ordering product",
    "independent ordering product",
    "independent product",
    "independently orderable",
    "separate upload row",
    "separate row",
    "listed separately",
    "this row is separated",
    "this row",
    "source product name",
    "keeps the source product name",
    "unified order model",
    "uniform order model",
    "visible source",
    "uploaded as a separate",
    "saved as a separate",
]
FORBIDDEN_PRODUCT_CONTAMINATION = [
    "battery",
    "batteries",
    "cathode",
    "anode",
    "precursor",
    "coin cell",
    "pouch cell",
    "energy storage",
    "solid-state battery",
    "battery materials",
    "cathode materials",
    "anode materials",
    "电池",
    "正极",
    "负极",
    "前驱体",
    "扣式",
    "软包",
    "储能",
]


def html_escape(value: Any) -> str:
    import html

    return html.escape(str(value or ""), quote=True)


def read_env_file(path: Path) -> None:
    if not path.exists():
        return
    for line in path.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if not stripped or stripped.startswith("#") or "=" not in stripped:
            continue
        key, value = stripped.split("=", 1)
        os.environ.setdefault(key.strip(), value.strip().strip('"').strip("'"))


def extract_docx_text(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table_index, table in enumerate(doc.tables, 1):
        parts.append(f"\n[TABLE {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_xlsx_text(path: Path) -> str:
    workbook = load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"\n[SHEET {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            values = [str(cell).strip() if cell is not None else "" for cell in row]
            if any(values):
                parts.append(" | ".join(values))
    return "\n".join(parts)


def extract_manual_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx_text(path)
    if suffix == ".xlsx":
        return extract_xlsx_text(path)
    if suffix in {".txt", ".md", ".csv"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Unsupported input type: {path}")


def cjk_exists(text: str) -> bool:
    return bool(re.search(r"[\u4e00-\u9fff]", text or ""))


def forbidden_phrase(text: str) -> str | None:
    lowered = (text or "").lower()
    for phrase in FORBIDDEN_PHRASES + FORBIDDEN_PRODUCT_CONTAMINATION:
        if phrase in lowered:
            return phrase
    return None


def clean_product_display_name(value: Any) -> str:
    text = str(value or "")
    text = re.sub(r"^\s*ATOMFAIR(?:®|Â®|Ã‚Â®|&reg;)?\s*", "", text, flags=re.I)
    return normalize_scientific_text(re.sub(r"\s+", " ", text).strip())


def normalize_scientific_text(value: Any) -> str:
    text = str(value or "")
    text = re.sub(r"\buL\b", "µL", text)
    text = re.sub(r"\bUL\b", "µL", text)
    text = text.replace("+/-", "±")
    text = re.sub(r"\bdeg\s*C\b", "°C", text, flags=re.I)
    text = re.sub(r"(?<=\d)\s*C\b", " °C", text)
    text = re.sub(r"(?<=\d)\s*℃\b", " °C", text)
    text = text.replace("Full high-temperature sterilization supported", "Fully autoclavable")
    text = text.replace("full high-temperature sterilization supported", "fully autoclavable")
    text = text.replace("Full autoclavable", "Fully autoclavable")
    text = text.replace("full autoclavable", "fully autoclavable")
    text = text.replace("Source product family", "Product Family")
    text = text.replace("Source Product Family", "Product Family")
    text = text.replace("source product family", "product family")
    text = text.replace("Source Model / Item", "Reference Model / Item")
    return text


def normalize_spec_value(spec: dict[str, Any]) -> dict[str, Any]:
    return {
        "parameter": normalize_scientific_text(spec.get("parameter", "")),
        "value": normalize_scientific_text(spec.get("value", "")),
    }


def split_test_volume_values(value: str) -> list[str]:
    text = normalize_scientific_text(value)
    parts = [part.strip() for part in re.split(r"\s*/\s*|;\s*", text) if part.strip()]
    volumes: list[str] = []
    for part in parts:
        if not re.search(r"\bµL\b", part):
            part = f"{part} µL"
        volumes.append(part)
    return volumes


def format_measurement_value(value: str) -> str:
    text = normalize_scientific_text(value.strip()).replace("+/-", "±")
    text = re.sub(r"^(systematic error|random error|accuracy|precision|inaccuracy|imprecision)\s*:\s*", "", text, flags=re.I)
    text = re.sub(r"(±[\d.]+\s*µL)\s+(±[\d.]+%)", r"\1 (\2)", text)
    return text


def pair_measurement_with_test_volumes(test_volumes: str, value: str) -> str:
    volumes = split_test_volume_values(test_volumes)
    text = normalize_scientific_text(value).replace("+/-", "±")
    text = re.sub(r"^(systematic error|random error|accuracy|precision|inaccuracy|imprecision)\s*:\s*", "", text.strip(), flags=re.I)
    parts = [format_measurement_value(part) for part in text.split(";") if part.strip()]
    if not volumes or len(volumes) != len(parts):
        return normalize_scientific_text(value)
    paired = []
    for index, (volume, part) in enumerate(zip(volumes, parts)):
        prefix = "At" if index == 0 else "at"
        paired.append(f"{prefix} {volume}: {part}")
    return "; ".join(paired)


def pair_measurement_specs(specs: list[dict[str, Any]]) -> list[dict[str, Any]]:
    test_volumes = ""
    for spec in specs:
        if str(spec.get("parameter", "")).strip().lower() in {"test volume(s)", "test volume", "test vol.", "test vol. (µl)"}:
            test_volumes = str(spec.get("value", ""))
            break
    if not test_volumes:
        return specs
    paired: list[dict[str, Any]] = []
    for spec in specs:
        parameter = str(spec.get("parameter", "")).strip()
        lower = parameter.lower()
        value = str(spec.get("value", ""))
        if any(term in lower for term in ["systematic error", "random error", "accuracy", "precision", "inaccuracy", "imprecision"]):
            spec = {"parameter": parameter, "value": pair_measurement_with_test_volumes(test_volumes, value)}
        paired.append(spec)
    return paired


def specs_with_atomfair_model(specs: list[dict[str, Any]], model: str) -> list[dict[str, Any]]:
    cleaned: list[dict[str, Any]] = []
    has_model = False
    for spec in specs or []:
        parameter = str(spec.get("parameter", "")).strip()
        if parameter.lower() == "atomfair model":
            cleaned.append({"parameter": "Atomfair Model", "value": model})
            has_model = True
        else:
            cleaned.append(normalize_spec_value(spec))
    if not has_model:
        cleaned.insert(0, {"parameter": "Atomfair Model", "value": model})
    return pair_measurement_specs(cleaned)


def safe_json_from_text(text: str) -> list[dict[str, Any]]:
    text = text.strip()
    fenced = re.search(r"```(?:json)?\s*(.*?)```", text, flags=re.S)
    if fenced:
        text = fenced.group(1).strip()
    start = text.find("[")
    end = text.rfind("]")
    if start >= 0 and end > start:
        text = text[start : end + 1]
    data = json.loads(text)
    if not isinstance(data, list):
        raise ValueError("AI response must be a JSON array.")
    return data


def build_prompt(source_text: str, brand: str, model_prefix: str) -> str:
    return f"""
You convert laboratory product manuals into upload-ready ecommerce records.

Return JSON only. Do not use markdown.

Non-negotiable rules:
- HTML and all English fields must contain English only.
- Keep Chinese product names from the source manual in chinese_name.
- Use source product names. If a source product name or description contains another manufacturer name, replace that manufacturer name with {brand}.
- Do not invent product names, product families, model meanings, ranges, capacities, dimensions, specifications, applications or compatibility.
- Do not manually generalize volume ranges, temperature ranges, capacities, speed ranges, accuracy tables or precision tables.
- Do not reuse a range table or specification table from one product family for another product family.
- Every distinct product specification, capacity, channel count, volume range, temperature range, rotor, block, module, configuration, accessory or spare part must become a separate item.
- Preserve 8-channel, 12-channel, 16-channel and all other multichannel variants as separate items.
- Accessories, spare parts, modules, electrodes, adapters, cables and consumables are products too; include every separately orderable item.
- Product Overview must be professional, detailed and source-supported. It may mention suitable laboratory or research applications when they fit the product and source evidence.
- Product Overview and HTML must describe the product page as complete data. Never mention internal workflow language such as "listed separately", "separate upload row", "unified order model", "source product name", "this row is separated", "independent ordering product", or similar phrasing.
- Backend category, tags and HTML source must only describe the current product page. Do not include unrelated product-family terms such as battery materials, cathode/anode materials, precursors, coin cells, pouch cells or energy-storage products unless the current source product is actually that product.
- Do not prefix english_name with ATOMFAIR, ATOMFAIR® or Atomfair. Keep brand identity in supplier/footer areas only.
- Put the Atomfair model in the Technical Specifications table as "Atomfair Model". Do not show "Atomfair Model" below the HTML title.
- Use the micro symbol `µL` for microliter units in product names, summaries and HTML. Do not write `uL` or `UL` in customer-facing text.
- Use professional sterilization wording such as "Fully autoclavable" or the source-supported cycle details. Do not write "Full high-temperature sterilization supported".
- Preserve source product casing such as HiPette-LTS. Do not force product names or HTML titles to all caps.
- Use "Product Family" instead of "Source Product Family" in customer-facing tables.
- ordering_rows is only for pages that truly compare multiple configurations in one HTML page. If the page represents a single specification, leave ordering_rows empty so the ordering table is not duplicated below Technical Specifications.
- Avoid repeating the same application sentence, selection guidance or parameter list across Product Overview, Key Features and Selection Note.
- If the source does not specify a value, write "Not specified" instead of guessing.
- source_column, upload_flag, alt_text, title and caption must always be empty strings.
- atomfair_model must be unique and start with {model_prefix}-.

Return a JSON array. Each item must use this shape:
{{
  "source_column": "",
  "chinese_name": "...",
  "atomfair_model": "{model_prefix}-...",
  "english_name": "...",
  "category": "...",
  "product_type": "...",
  "overview_paragraphs": ["...", "..."],
  "features": ["...", "..."],
  "specifications": [{{"parameter": "...", "value": "..."}}],
  "accessories": [{{"name": "...", "reference_code": "", "specification": "...", "compatibility": "..."}}],
  "ordering_rows": [{{"order_model": "...", "source_model": "...", "key_specification": "...", "application": "..."}}],
  "source_evidence": ["source heading or table reference used for this item"],
  "upload_flag": "",
  "alt_text": "",
  "title": "",
  "caption": ""
}}

Manual text:
{source_text}
""".strip()


def call_openai(prompt: str, model: str) -> str:
    try:
        from openai import OpenAI
    except Exception as exc:
        raise RuntimeError("OpenAI package is not installed. Run: pip install -r requirements.txt") from exc
    if not os.environ.get("OPENAI_API_KEY"):
        raise RuntimeError("OPENAI_API_KEY is missing. Copy .env.example to .env and fill in your key.")
    client = OpenAI()
    response = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=0,
    )
    return response.choices[0].message.content or ""


def html_table(headers: list[str], rows: list[list[str]]) -> str:
    head_cells = "".join(
        f'<th style="background:#111;color:#fff;border:1px solid #111;padding:11px 12px;text-align:left;">{html_escape(header)}</th>'
        for header in headers
    )
    body_rows: list[str] = []
    for row in rows:
        cells = []
        for cell in row:
            cells.append(
                f'<td style="border:1px solid #d7d7d7;padding:10px 12px;vertical-align:top;text-align:left;">{html_escape(cell)}</td>'
            )
        body_rows.append("<tr>" + "".join(cells) + "</tr>")
    return (
        '<table style="border-collapse:collapse;border:1px solid #111;font-size:14px;" border="0" '
        f'width="100%" cellspacing="0" cellpadding="0"><thead><tr>{head_cells}</tr></thead>'
        f'<tbody>{"".join(body_rows)}</tbody></table>'
    )


def html_section(title: str, body: str) -> str:
    return (
        '<table style="margin-bottom:34px;" border="0" width="100%" cellspacing="0" cellpadding="0"><tbody>'
        '<tr><td style="border-left:5px solid #111;padding-left:15px;padding-bottom:12px;">'
        f'<h2 style="margin:0;font-size:18px;color:#111;text-transform:uppercase;font-weight:800;">{html_escape(title)}</h2>'
        f'</td></tr><tr><td style="padding-top:18px;">{body}</td></tr></tbody></table>'
    )


def product_html(item: dict[str, Any], brand: str) -> str:
    overview = [normalize_scientific_text(text) for text in (item.get("overview_paragraphs") or [])]
    features = [normalize_scientific_text(text) for text in (item.get("features") or [])]
    english_name = clean_product_display_name(item.get("english_name"))
    model = str(item.get("atomfair_model", "")).strip()
    specs = specs_with_atomfair_model(item.get("specifications") or [], model)
    accessories = item.get("accessories") or []
    ordering = item.get("ordering_rows") or []

    overview_html = "".join(
        f'<p style="margin:0 0 14px 0;font-size:15px;color:#444;text-align:justify;">{html_escape(paragraph)}</p>'
        for paragraph in overview
    )
    feature_html = (
        '<ul style="margin:0;padding-left:22px;font-size:14px;color:#444;line-height:1.85;">'
        + "".join(f'<li style="margin-bottom:9px;">{html_escape(feature)}</li>' for feature in features)
        + "</ul>"
    )
    spec_html = html_table(
        ["Parameter", "Specification / Available Values"],
        [[spec.get("parameter", ""), spec.get("value", "")] for spec in specs],
    )

    parts = [
        '<!-- Atomfair Complete Product HTML - Full-English Release -->',
        '<div style="width:100%;background:#ffffff;padding:0;" align="center">',
        '<table style="width:100%;font-family:Arial,Helvetica,sans-serif;color:#333;line-height:1.58;background:#fff;" border="0" width="100%" cellspacing="0" cellpadding="0"><tbody>',
        '<tr><td style="padding:38px 20px;border-bottom:4px solid #111;">',
        f'<h1 style="margin:0;font-size:26px;color:#111;font-weight:800;">{html_escape(english_name)}</h1>',
        f'<div style="margin-top:8px;font-size:14px;color:#333;">Product Type: {html_escape(item.get("product_type"))}</div>',
        '<div style="margin-top:15px;display:inline-block;background:#111;color:#fff;padding:6px 15px;font-size:13px;font-weight:bold;letter-spacing:.6px;text-transform:uppercase;">Research-grade laboratory product</div>',
        '</td></tr><tr><td style="padding:38px 20px;">',
        html_section("Product Overview", overview_html),
        html_section("Key Features and Advantages", feature_html),
        html_section("Technical Specifications", spec_html),
    ]
    if accessories:
        parts.append(
            html_section(
                "Compatible Accessories and Configuration Data",
                html_table(
                    ["Accessory / Module", "Reference Code", "Specification", "Compatibility / Use"],
                    [
                        [
                            accessory.get("name", ""),
                            accessory.get("reference_code", ""),
                            normalize_scientific_text(accessory.get("specification", "")),
                            normalize_scientific_text(accessory.get("compatibility", "")),
                        ]
                        for accessory in accessories
                    ],
                ),
            )
        )
    if len(ordering) > 1:
        parts.append(
            html_section(
                "Atomfair Product Ordering and Configuration Table",
                html_table(
                    ["Unique Atomfair Order Model", "Source Model / Item", "Key Specification", "Primary Application"],
                    [
                        [
                            row.get("order_model", ""),
                            normalize_scientific_text(row.get("source_model", "")),
                            normalize_scientific_text(row.get("key_specification", "")),
                            normalize_scientific_text(row.get("application", "")),
                        ]
                        for row in ordering
                    ],
                ),
            )
        )
    parts.extend(
        [
            '<table style="background:#f7f7f7;border:1px solid #e0e0e0;margin-bottom:34px;" border="0" width="100%" cellspacing="0" cellpadding="22"><tbody><tr><td style="font-size:14px;color:#333;">',
            '<div style="margin-bottom:10px;"><strong>QUALITY ASSURANCE:</strong> Product configuration, specification matching and technical documentation are reviewed for research and laboratory procurement.</div>',
            '<div style="margin-bottom:10px;"><strong>AFTER-SALES SUPPORT:</strong> Accessory matching, replacement part support and configuration confirmation are available for laboratory users.</div>',
            '<div style="margin-bottom:10px;"><strong>COMPATIBLE CONSUMABLES:</strong> Compatible consumables and application-specific accessories can be recommended according to the selected product configuration.</div>',
            '<div style="border:1px solid #f0c75e;background:#fff8df;padding:14px 16px;margin-top:16px;"><strong>HOW TO PURCHASE:</strong> Contact the Atomfair sales team at inquiry@atomfair.com with the required model, configuration and quantity for quotation support.</div>',
            '</td></tr></tbody></table></td></tr>',
            '<tr><td style="padding:25px 20px;background:#111;text-align:center;">',
            '<div style="color:#fff;font-size:20px;font-weight:bold;margin-bottom:8px;letter-spacing:1px;">TAILORED SOLUTIONS FOR RESEARCH</div>',
            '<div style="color:#ddd;font-size:13px;margin-bottom:18px;">For bulk orders, OEM requirements, accessory matching, platform selection or configuration confirmation, contact our engineering sales team.</div>',
            '<div style="display:inline-block;background:#fff;color:#111;padding:12px 35px;font-weight:800;font-size:13px;text-transform:uppercase;letter-spacing:1px;border-radius:2px;">EMAIL: inquiry@atomfair.com</div>',
            '</td></tr>',
            '<tr><td style="padding:28px 20px;text-align:center;font-size:12px;color:#888;letter-spacing:1px;">',
            f'<div style="margin-bottom:5px;text-transform:uppercase;"><strong>Supplier:</strong> {html_escape(brand)}</div>',
            f'<div style="text-transform:uppercase;"><strong>Brand:</strong> {html_escape(brand.upper())}&reg;</div>',
            '</td></tr></tbody></table></div>',
        ]
    )
    return "".join(parts)


def validate_items(items: list[dict[str, Any]], model_prefix: str) -> None:
    seen_models: set[str] = set()
    required = [
        "chinese_name",
        "atomfair_model",
        "english_name",
        "category",
        "product_type",
        "overview_paragraphs",
        "features",
        "specifications",
    ]
    for index, item in enumerate(items, 1):
        missing = [field for field in required if not item.get(field)]
        if missing:
            raise ValueError(f"Item {index} is missing required fields: {', '.join(missing)}")
        model = str(item.get("atomfair_model", "")).strip()
        if not model.startswith(f"{model_prefix}-"):
            raise ValueError(f"Item {index} model must start with {model_prefix}-: {model}")
        if model in seen_models:
            raise ValueError(f"Duplicate model generated: {model}")
        seen_models.add(model)
        english_text = "\n".join(
            [
                clean_product_display_name(item.get("english_name", "")),
                str(item.get("category", "")),
                str(item.get("product_type", "")),
                "\n".join(map(str, item.get("overview_paragraphs") or [])),
                "\n".join(map(str, item.get("features") or [])),
            ]
        )
        if cjk_exists(english_text):
            raise ValueError(f"English fields contain Chinese characters for model {model}")
        phrase = forbidden_phrase(english_text)
        if phrase:
            raise ValueError(f"Forbidden workflow phrase found for model {model}: {phrase}")


def make_records(items: list[dict[str, Any]], brand: str) -> list[dict[str, str]]:
    records: list[dict[str, str]] = []
    for item in items:
        item["english_name"] = clean_product_display_name(item.get("english_name"))
        code = product_html(item, brand)
        if cjk_exists(code):
            raise ValueError(f"HTML contains Chinese characters for model {item.get('atomfair_model')}.")
        phrase = forbidden_phrase(code)
        if phrase:
            raise ValueError(f"Forbidden workflow phrase found in HTML for model {item.get('atomfair_model')}: {phrase}")

        row = {header: "" for header in HEADERS}
        row.update(
            {
                "商品中文名称": normalize_scientific_text(item.get("chinese_name", "")),
                "型号": str(item.get("atomfair_model", "")),
                "产品名称（英文）": clean_product_display_name(item.get("english_name", "")),
                "简介": normalize_scientific_text(" ".join(map(str, item.get("overview_paragraphs") or []))),
                "代码": code,
                "网站分类": str(item.get("category", "")),
            }
        )
        for column in BLANK_UPLOAD_COLUMNS:
            row[column] = ""
        for english_column in ["产品名称（英文）", "简介", "代码", "网站分类"]:
            if cjk_exists(row.get(english_column, "")):
                raise ValueError(f"{english_column} contains Chinese characters for model {item.get('atomfair_model')}.")
        phrase = forbidden_phrase(row["简介"])
        if phrase:
            raise ValueError(f"Forbidden workflow phrase found in overview for model {item.get('atomfair_model')}: {phrase}")
        records.append(row)
    return records


def write_excel(records: list[dict[str, str]], output_path: Path, sheet_name: str, style_reference: Path | None, style_sheet: str) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = sheet_name[:31] or "Products"
    sheet.append(HEADERS)
    for record in records:
        sheet.append([record.get(header, "") for header in HEADERS])

    if style_reference and style_reference.exists():
        apply_reference_style(sheet, style_reference, style_sheet)
    else:
        apply_default_upload_style(sheet)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(output_path)


def apply_default_upload_style(sheet) -> None:
    blue_fill = PatternFill("solid", fgColor="C0E6F5")
    white_fill = PatternFill("solid", fgColor="FFFFFF")
    font = Font(name="Calibri", size=11, color="000000")
    thin = Side(style="thin", color="D9D9D9")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    widths = [12, 24, 22, 36, 16, 66, 18, 90, 22, 14, 14, 18, 20, 18, 18, 12, 8, 12]
    for index, width in enumerate(widths, 1):
        sheet.column_dimensions[sheet.cell(1, index).column_letter].width = width
    for row_index, row in enumerate(sheet.iter_rows(), 1):
        fill = blue_fill if row_index % 2 == 0 or row_index == 1 else white_fill
        for cell in row:
            cell.border = border
            cell.fill = fill
            cell.font = font
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    for cell in sheet[1]:
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    sheet.freeze_panes = None
    table = Table(displayName="UploadTable", ref=f"A1:R{sheet.max_row}")
    table.tableStyleInfo = TableStyleInfo(
        name="TableStyleMedium2",
        showFirstColumn=False,
        showLastColumn=False,
        showRowStripes=True,
        showColumnStripes=False,
    )
    sheet.add_table(table)


def apply_reference_style(sheet, reference_path: Path, style_sheet: str) -> None:
    from copy import copy

    reference_wb = load_workbook(reference_path)
    reference_ws = reference_wb[style_sheet] if style_sheet in reference_wb.sheetnames else reference_wb.active
    max_col = min(sheet.max_column, reference_ws.max_column)
    for col_index in range(1, sheet.max_column + 1):
        letter = sheet.cell(1, col_index).column_letter
        ref_dim = reference_ws.column_dimensions[letter]
        sheet.column_dimensions[letter].width = ref_dim.width
        sheet.column_dimensions[letter].hidden = ref_dim.hidden
    for row_index in range(1, sheet.max_row + 1):
        ref_row_index = row_index if row_index <= reference_ws.max_row else ((row_index - 2) % 2) + 2
        sheet.row_dimensions[row_index].height = reference_ws.row_dimensions[ref_row_index].height
        for col_index in range(1, sheet.max_column + 1):
            ref_cell = reference_ws.cell(ref_row_index, col_index if col_index <= max_col else max_col)
            cell = sheet.cell(row_index, col_index)
            cell.font = copy(ref_cell.font)
            cell.fill = copy(ref_cell.fill)
            cell.border = copy(ref_cell.border)
            cell.alignment = copy(ref_cell.alignment)
            cell.number_format = ref_cell.number_format
            cell.protection = copy(ref_cell.protection)
    sheet.freeze_panes = reference_ws.freeze_panes
    sheet.sheet_view.showGridLines = reference_ws.sheet_view.showGridLines
    if reference_ws.tables:
        ref_table = next(iter(reference_ws.tables.values()))
        table = Table(displayName="UploadTable", ref=f"A1:R{sheet.max_row}")
        if ref_table.tableStyleInfo:
            table.tableStyleInfo = TableStyleInfo(
                name=ref_table.tableStyleInfo.name,
                showFirstColumn=ref_table.tableStyleInfo.showFirstColumn,
                showLastColumn=ref_table.tableStyleInfo.showLastColumn,
                showRowStripes=ref_table.tableStyleInfo.showRowStripes,
                showColumnStripes=ref_table.tableStyleInfo.showColumnStripes,
            )
        sheet.add_table(table)


def validate_records(records: list[dict[str, str]]) -> None:
    models = [record["型号"] for record in records]
    if len(models) != len(set(models)):
        raise ValueError("Duplicate models found after record conversion.")
    for record in records:
        for column in BLANK_UPLOAD_COLUMNS:
            if record.get(column):
                raise ValueError(f"Column {column} must be blank for model {record.get('型号')}.")
        if cjk_exists(record.get("代码", "")):
            raise ValueError(f"HTML contains Chinese characters for model {record.get('型号')}.")
        phrase = forbidden_phrase(record.get("代码", "") + "\n" + record.get("简介", ""))
        if phrase:
            raise ValueError(f"Forbidden workflow phrase found for model {record.get('型号')}: {phrase}")


def collect_input_files(input_path: Path) -> list[Path]:
    if input_path.is_dir():
        return sorted(path for path in input_path.iterdir() if path.suffix.lower() in SUPPORTED_EXTENSIONS and not path.name.startswith("~$"))
    return [input_path]


def main() -> None:
    read_env_file(Path(".env"))
    parser = argparse.ArgumentParser(description="Generate full-English product HTML and upload-ready Excel from product manuals.")
    parser.add_argument("--input", "-i", default="input", help="Manual file or folder. Supports .docx, .xlsx, .txt, .md and .csv.")
    parser.add_argument("--output", "-o", default="outputs/products.xlsx", help="Output XLSX path.")
    parser.add_argument("--records", default="outputs/products_records.json", help="Output JSON records path.")
    parser.add_argument("--sheet-name", default="Products", help="Excel sheet name.")
    parser.add_argument("--brand", default=os.environ.get("BRAND_NAME", "Atomfair"), help="Brand name used in generated HTML.")
    parser.add_argument("--model-prefix", default=os.environ.get("MODEL_PREFIX", "AF"), help="Required model prefix.")
    parser.add_argument("--model", default=os.environ.get("OPENAI_MODEL", "gpt-4.1-mini"), help="OpenAI model name.")
    parser.add_argument("--style-reference", default="", help="Optional workbook whose style should be copied, for example a Pipetting Series upload file.")
    parser.add_argument("--style-sheet", default="Pipetting Series", help="Sheet name to copy from --style-reference.")
    args = parser.parse_args()

    input_files = collect_input_files(Path(args.input))
    if not input_files:
        raise SystemExit("No supported input files found. Put manuals into input/ or pass --input path/to/manual.docx")

    all_items: list[dict[str, Any]] = []
    for input_file in input_files:
        source_text = extract_manual_text(input_file)
        if not source_text.strip():
            raise SystemExit(f"No readable text or tables found in {input_file}. Use Codex visual/OCR review for image-only manuals.")
        prompt = build_prompt(source_text, args.brand, args.model_prefix)
        ai_text = call_openai(prompt, args.model)
        all_items.extend(safe_json_from_text(ai_text))

    validate_items(all_items, args.model_prefix)
    records = make_records(all_items, args.brand)
    validate_records(records)

    records_path = Path(args.records)
    records_path.parent.mkdir(parents=True, exist_ok=True)
    records_path.write_text(json.dumps([{"sheetName": args.sheet_name, "records": records}], ensure_ascii=False, indent=2), encoding="utf-8")
    style_reference = Path(args.style_reference) if args.style_reference else None
    write_excel(records, Path(args.output), args.sheet_name, style_reference, args.style_sheet)

    print(
        json.dumps(
            {
                "input_files": [str(path) for path in input_files],
                "items": len(records),
                "records": str(records_path),
                "excel": args.output,
                "blank_columns": sorted(BLANK_UPLOAD_COLUMNS),
                "html_contains_chinese": False,
                "style_reference": str(style_reference) if style_reference else "",
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
