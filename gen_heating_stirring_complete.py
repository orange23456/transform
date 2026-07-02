from pathlib import Path
from docx import Document
from docx.shared import Pt
import html as htmlmod
import re
import sys
import zipfile

try:
    sys.stdout.reconfigure(encoding='utf-8')
except Exception:
    pass

SOURCE = Path(r'C:\Users\hz-user\Documents\HTML生成\gen_heating_stirring.py')
OUT = Path(r'C:\Users\hz-user\Desktop\全自动\generated_html_full\Atomfair_heating_stirring_HTML')
OUT.mkdir(parents=True, exist_ok=True)

source = SOURCE.read_text(encoding='utf-8-sig', errors='ignore')
marker = '\nfor p in products:'
if marker not in source:
    raise SystemExit('Cannot find product data marker')
namespace = {'__name__': '__not_main__'}
exec(compile(source.split(marker, 1)[0], str(SOURCE), 'exec'), namespace)
products = namespace['products']


def esc(value):
    return htmlmod.escape(str(value or ''), quote=True)


def cn_name(p):
    return Path(p['file']).stem

def html_product_name(p):
    return str(p['title'])


def base_model(p):
    # Keep the unified Atomfair model family already assigned by our rule set.
    return re.sub(r'[^A-Z0-9-]', '', str(p['model']).upper()) or 'AF-HS'


def rows_html(specs):
    return '\n'.join(
        f'<tr><td style="border:1px solid #d7d7d7;padding:10px 12px;vertical-align:top;background:#ffffff;font-weight:700;color:#222;">{esc(k)}</td><td style="border:1px solid #d7d7d7;padding:10px 12px;vertical-align:top;background:#ffffff;color:#333;">{esc(v)}</td></tr>'
        for k, v in specs
    )


def bullets_html(items):
    return '\n'.join(f'<li style="margin-bottom:9px;">{esc(x)}</li>' for x in items)


def descriptive_paragraphs(p):
    name = html_product_name(p)
    return [
        f"{name} is prepared under the Atomfair brand as a research-grade laboratory product. The product category and technical meaning follow the source manual, while supplier and brand identity are presented as Atomfair.",
        p['overview'],
        f"This product is suitable for {p['apps'].rstrip('.')}. The narrative description explains use and positioning only; technical data, dimensions, ranges, capacities and configuration differences are kept in structured tables below so that no available specification is hidden or compressed.",
    ]


def split_values(value):
    text = str(value).strip()
    # Split only where the text clearly lists alternatives; keep explanatory sentences intact.
    if ';' in text:
        parts = [x.strip() for x in text.split(';') if x.strip()]
    elif ' or ' in text:
        parts = [x.strip() for x in re.split(r'\s+or\s+', text) if x.strip()]
    elif re.search(r'\bdepending on\b|\bconfigurations?\b|\bformats?\b', text, re.I) and ',' in text:
        parts = [x.strip() for x in text.split(',') if x.strip()]
    else:
        parts = [text]
    return parts


def configuration_rows(p):
    rows = []
    base = base_model(p)
    source_fields = []
    for key, value in p['specs']:
        value_text = str(value)
        is_variant_field = (
            'depending on' in value_text.lower()
            or 'configuration' in value_text.lower()
            or 'configurations' in value_text.lower()
            or 'formats' in value_text.lower()
            or ' or ' in value_text
            or ('visible' in value_text.lower() and ',' in value_text)
        )
        if is_variant_field:
            source_fields.append((key, split_values(value_text)))

    if source_fields:
        idx = 1
        for key, parts in source_fields:
            for part in parts:
                rows.append((f'{base}-{idx:02d}', html_product_name(p), key, part))
                idx += 1
    else:
        rows.append((f'{base}-01', html_product_name(p), 'Standard configuration', 'Single available configuration based on the source manual data.'))
    return rows


def configuration_rows_html(rows):
    return '\n'.join(
        f'<tr><td style="border:1px solid #d7d7d7;padding:10px 12px;background:#ffffff;font-weight:700;">{esc(model)}</td><td style="border:1px solid #d7d7d7;padding:10px 12px;background:#ffffff;">{esc(name)}</td><td style="border:1px solid #d7d7d7;padding:10px 12px;background:#ffffff;">{esc(field)}</td><td style="border:1px solid #d7d7d7;padding:10px 12px;background:#ffffff;">{esc(detail)}</td></tr>'
        for model, name, field, detail in rows
    )


def paragraphs_html(paragraphs):
    return '\n'.join(f'<p style="margin:0 0 14px 0;font-size:15px;color:#444444;text-align:justify;">{esc(text)}</p>' for text in paragraphs)


def make_full_html(p):
    name = html_product_name(p)
    base = base_model(p)
    overview_paragraphs = descriptive_paragraphs(p)
    config_rows = configuration_rows(p)
    return f'''<!-- Atomfair Complete Product HTML - Source-Name + Unified-Model Release -->
<div style="width:100%;background:#ffffff;padding:0;" align="center">
<table style="width:100%;max-width:100%;background:#ffffff;font-family:Arial,Helvetica,sans-serif;color:#333333;line-height:1.58;" border="0" width="100%" cellspacing="0" cellpadding="0">
<tbody>
<tr><td style="padding:38px 20px;border-bottom:4px solid #111111;">
<h1 style="margin:0;font-size:26px;color:#111111;font-weight:800;text-transform:uppercase;">ATOMFAIR&reg; {esc(name)}</h1>
<div style="margin-top:8px;font-size:14px;color:#333333;">Product Name: {esc(name)}</div>
<div class="model" style="margin-top:10px;font-size:15px;color:#333333;font-weight:700;">Atomfair Model: {esc(base)} Series</div>
<div style="margin-top:15px;display:inline-block;background:#111111;color:#ffffff;padding:6px 15px;font-size:13px;font-weight:bold;letter-spacing:.6px;text-transform:uppercase;">Research-grade stirring and heating equipment for laboratory workflows</div>
</td></tr>
<tr><td style="padding:38px 20px;">
<table style="margin-bottom:34px;" border="0" width="100%" cellspacing="0" cellpadding="0"><tbody>
<tr><td style="border-left:5px solid #111111;padding-left:15px;padding-bottom:12px;"><h2 style="margin:0;font-size:18px;color:#111111;text-transform:uppercase;font-weight:800;">Product Overview</h2></td></tr>
<tr><td style="padding-top:18px;">{paragraphs_html(overview_paragraphs)}</td></tr>
</tbody></table>
<table style="margin-bottom:34px;" border="0" width="100%" cellspacing="0" cellpadding="0"><tbody>
<tr><td style="border-left:5px solid #111111;padding-left:15px;padding-bottom:12px;"><h2 style="margin:0;font-size:18px;color:#111111;text-transform:uppercase;font-weight:800;">Key Features and Advantages</h2></td></tr>
<tr><td><ul style="margin:18px 0 0 0;padding-left:22px;font-size:14px;color:#444444;line-height:1.85;">{bullets_html(p['features'])}</ul></td></tr>
</tbody></table>
<table style="margin-bottom:34px;" border="0" width="100%" cellspacing="0" cellpadding="0"><tbody>
<tr><td style="border-left:5px solid #111111;padding-left:15px;padding-bottom:12px;"><h2 style="margin:0;font-size:18px;color:#111111;text-transform:uppercase;font-weight:800;">Technical Specifications</h2></td></tr>
<tr><td><table style="border-collapse:collapse;border:1px solid #111111;font-size:14px;" border="0" width="100%" cellspacing="0" cellpadding="0">
<thead><tr><th style="background:#111111;color:#ffffff;border:1px solid #111111;padding:11px 12px;text-align:left;">Parameter</th><th style="background:#111111;color:#ffffff;border:1px solid #111111;padding:11px 12px;text-align:left;">Specification / Available Values</th></tr></thead>
<tbody>
{rows_html(p['specs'])}
</tbody></table></td></tr>
</tbody></table>
<table style="margin-bottom:34px;" border="0" width="100%" cellspacing="0" cellpadding="0"><tbody>
<tr><td style="border-left:5px solid #111111;padding-left:15px;padding-bottom:12px;"><h2 style="margin:0;font-size:18px;color:#111111;text-transform:uppercase;font-weight:800;">Atomfair Product Ordering and Configuration Table</h2></td></tr>
<tr><td><table style="border-collapse:collapse;border:1px solid #111111;font-size:14px;" border="0" width="100%" cellspacing="0" cellpadding="0">
<thead><tr><th style="background:#111111;color:#ffffff;border:1px solid #111111;padding:11px 12px;text-align:left;">Unique Atomfair Order Model</th><th style="background:#111111;color:#ffffff;border:1px solid #111111;padding:11px 12px;text-align:left;">Manual Product Name</th><th style="background:#111111;color:#ffffff;border:1px solid #111111;padding:11px 12px;text-align:left;">Differentiating Parameter</th><th style="background:#111111;color:#ffffff;border:1px solid #111111;padding:11px 12px;text-align:left;">Specification Value</th></tr></thead>
<tbody>
{configuration_rows_html(config_rows)}
</tbody></table></td></tr>
</tbody></table>
<table style="background:#f7f7f7;border:1px solid #e0e0e0;margin-bottom:34px;" border="0" width="100%" cellspacing="0" cellpadding="22"><tbody>
<tr><td style="font-size:14px;color:#333333;"><div style="margin-bottom:10px;"><strong>Application Scope:</strong> {esc(p['apps'])}</div><div style="margin-bottom:10px;"><strong>Naming Standard:</strong> Product names are presented in English only. Supplier or brand identity is shown as Atomfair where required. Atomfair model numbers are assigned separately using a unified order-model rule.</div><div><strong>Ordering Standard:</strong> Each different specification or configuration receives its own unique Atomfair order model in the ordering table.</div></td></tr>
</tbody></table>
</td></tr>
<tr><td style="padding:25px 20px;background:#111111;text-align:center;"><div style="color:#ffffff;font-size:20px;font-weight:bold;margin-bottom:8px;letter-spacing:1px;">TAILORED SOLUTIONS FOR RESEARCH</div><div style="color:#dddddd;font-size:13px;margin-bottom:18px;">For bulk orders, OEM requirements, or configuration matching, contact our engineering sales team.</div><div style="display:inline-block;background:#ffffff;color:#111111;padding:12px 35px;font-weight:800;font-size:13px;text-transform:uppercase;letter-spacing:1px;border-radius:2px;">EMAIL: inquiry@atomfair.com</div></td></tr>
<tr><td style="padding:28px 20px;text-align:center;font-size:12px;color:#888888;letter-spacing:1px;"><div style="margin-bottom:5px;text-transform:uppercase;"><strong>Supplier:</strong> Atomfair</div><div style="text-transform:uppercase;"><strong>Brand:</strong> ATOMFAIR&reg;</div></td></tr>
</tbody></table></div>'''

for p in products:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(36)
    sec.bottom_margin = Pt(36)
    sec.left_margin = Pt(36)
    sec.right_margin = Pt(36)
    style = doc.styles['Normal']
    style.font.name = 'Consolas'
    style.font.size = Pt(8.5)
    html = make_full_html(p)
    for line in html.splitlines():
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
    doc.save(OUT / p['file'])

failures = []
for p in products:
    path = OUT / p['file']
    with zipfile.ZipFile(path) as z:
        xml = z.read('word/document.xml').decode('utf-8', errors='replace')
    text = re.sub(r'<[^>]+>', '', xml)
    ok = path.exists() and html_product_name(p) in text and 'Unique Atomfair Order Model' in text and base_model(p) in text and 'TAILORED SOLUTIONS FOR RESEARCH' in text and not re.search(r'[\u4e00-\u9fff]', text)
    print(path.name + ' OK=' + str(ok))
    if not ok:
        failures.append(path.name)
if failures:
    raise SystemExit('QC failures: ' + ', '.join(failures))
print('OUTPUT_DIR=' + str(OUT))
print('COUNT=' + str(len(products)))
